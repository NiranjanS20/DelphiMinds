"""
Advanced PDF and document processing utilities for resume analysis
"""
import io
import tempfile
from typing import Dict, Any, Optional, Tuple
from pathlib import Path

# Import with fallbacks
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class AdvancedDocumentProcessor:
    """Advanced document processing with multiple extraction methods"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def extract_text_from_file(self, file_obj, filename: str) -> Dict[str, Any]:
        """
        Extract text and metadata from various document formats
        Returns a dictionary with text, metadata, and extraction method used
        """
        file_extension = Path(filename).suffix.lower()
        
        result = {
            'text': '',
            'metadata': {},
            'extraction_method': 'unknown',
            'success': False,
            'error': None
        }
        
        try:
            if file_extension == '.pdf':
                result = self._extract_from_pdf(file_obj)
            elif file_extension in ['.docx', '.doc']:
                result = self._extract_from_docx(file_obj)
            elif file_extension == '.txt':
                result = self._extract_from_txt(file_obj)
            else:
                result['error'] = f"Unsupported file format: {file_extension}"
                
        except Exception as e:
            result['error'] = f"Error processing file: {str(e)}"
        
        return result
    
    def _extract_from_pdf(self, file_obj) -> Dict[str, Any]:
        """Extract text from PDF using multiple methods as fallbacks"""
        methods_to_try = []
        
        if PDFPLUMBER_AVAILABLE:
            methods_to_try.append(('pdfplumber', self._extract_with_pdfplumber))
        if FITZ_AVAILABLE:
            methods_to_try.append(('pymupdf', self._extract_with_pymupdf))
        if PYPDF2_AVAILABLE:
            methods_to_try.append(('pypdf2', self._extract_with_pypdf2))
        
        for method_name, method_func in methods_to_try:
            try:
                file_obj.seek(0)  # Reset file pointer
                result = method_func(file_obj)
                if result['success'] and result['text'].strip():
                    result['extraction_method'] = method_name
                    return result
            except Exception as e:
                continue
        
        return {
            'text': '',
            'metadata': {},
            'extraction_method': 'none',
            'success': False,
            'error': 'All PDF extraction methods failed'
        }
    
    def _extract_with_pdfplumber(self, file_obj) -> Dict[str, Any]:
        """Extract text using pdfplumber (best for tables and complex layouts)"""
        text = ""
        metadata = {'pages': 0, 'tables': 0}
        
        with pdfplumber.open(file_obj) as pdf:
            metadata['pages'] = len(pdf.pages)
            
            for page in pdf.pages:
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                
                # Extract tables
                tables = page.extract_tables()
                if tables:
                    metadata['tables'] += len(tables)
                    for table in tables:
                        # Convert table to text
                        for row in table:
                            if row:
                                text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
        
        return {
            'text': text,
            'metadata': metadata,
            'success': True,
            'error': None
        }
    
    def _extract_with_pymupdf(self, file_obj) -> Dict[str, Any]:
        """Extract text using PyMuPDF (good for general text extraction)"""
        text = ""
        metadata = {'pages': 0}
        
        doc = fitz.open(stream=file_obj.read(), filetype="pdf")
        metadata['pages'] = len(doc)
        
        for page in doc:
            text += page.get_text() + "\n"
        
        doc.close()
        
        return {
            'text': text,
            'metadata': metadata,
            'success': True,
            'error': None
        }
    
    def _extract_with_pypdf2(self, file_obj) -> Dict[str, Any]:
        """Extract text using PyPDF2 (fallback option)"""
        text = ""
        metadata = {'pages': 0}
        
        pdf_reader = PyPDF2.PdfReader(file_obj)
        metadata['pages'] = len(pdf_reader.pages)
        
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return {
            'text': text,
            'metadata': metadata,
            'success': True,
            'error': None
        }
    
    def _extract_from_docx(self, file_obj) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            return {
                'text': '',
                'metadata': {},
                'extraction_method': 'none',
                'success': False,
                'error': 'python-docx not available'
            }
        
        try:
            doc = docx.Document(file_obj)
            text = ""
            metadata = {'paragraphs': 0, 'tables': 0}
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                if paragraph.text.strip():
                    metadata['paragraphs'] += 1
            
            # Extract tables
            for table in doc.tables:
                metadata['tables'] += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += " | ".join(row_text) + "\n"
            
            return {
                'text': text,
                'metadata': metadata,
                'extraction_method': 'python-docx',
                'success': True,
                'error': None
            }
            
        except Exception as e:
            return {
                'text': '',
                'metadata': {},
                'extraction_method': 'none',
                'success': False,
                'error': f"DOCX extraction failed: {str(e)}"
            }
    
    def _extract_from_txt(self, file_obj) -> Dict[str, Any]:
        """Extract text from plain text files"""
        try:
            file_obj.seek(0)
            text = file_obj.read().decode('utf-8')
            
            return {
                'text': text,
                'metadata': {'lines': len(text.split('\n'))},
                'extraction_method': 'plain_text',
                'success': True,
                'error': None
            }
            
        except Exception as e:
            return {
                'text': '',
                'metadata': {},
                'extraction_method': 'none',
                'success': False,
                'error': f"Text extraction failed: {str(e)}"
            }
    
    def analyze_document_quality(self, extraction_result: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze the quality of extracted text"""
        text = extraction_result.get('text', '')
        
        if not text.strip():
            return {
                'quality_score': 0,
                'issues': ['No text extracted'],
                'suggestions': ['Try a different file format or scan quality']
            }
        
        quality_score = 100
        issues = []
        suggestions = []
        
        # Check text length
        if len(text) < 100:
            quality_score -= 30
            issues.append('Very short text content')
            suggestions.append('Ensure the document contains substantial content')
        
        # Check for garbled text (too many special characters)
        special_char_ratio = len([c for c in text if not c.isalnum() and not c.isspace()]) / len(text)
        if special_char_ratio > 0.3:
            quality_score -= 20
            issues.append('High ratio of special characters (possible OCR issues)')
            suggestions.append('Consider using a higher quality scan or different PDF')
        
        # Check for repeated characters (OCR artifacts)
        repeated_chars = len([c for c in text if text.count(c) > len(text) * 0.1])
        if repeated_chars > 0:
            quality_score -= 15
            issues.append('Repeated character patterns detected')
            suggestions.append('Document may have OCR artifacts')
        
        # Check for proper sentence structure
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        avg_sentence_length = sum(len(s.split()) for s in sentences) / max(len(sentences), 1)
        
        if avg_sentence_length < 3:
            quality_score -= 25
            issues.append('Very short average sentence length')
            suggestions.append('Text may be fragmented or poorly extracted')
        
        return {
            'quality_score': max(0, quality_score),
            'issues': issues,
            'suggestions': suggestions,
            'text_stats': {
                'length': len(text),
                'words': len(text.split()),
                'sentences': len(sentences),
                'avg_sentence_length': round(avg_sentence_length, 1)
            }
        }